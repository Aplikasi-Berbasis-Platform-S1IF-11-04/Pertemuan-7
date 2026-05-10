import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('LAPORAN PRAKTIKUM\nAPLIKASI BERBASIS PLATFORM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('MODUL - 12 & 13\nLARAVEL: DATABASE 2 (AUTH, MIDDLEWARE & RELATIONS) & GIT BRANCHING', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Author
    auth_heading = doc.add_heading('Disusun Oleh :', 2)
    auth_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth = doc.add_paragraph('Haposan Felix Marcel Siregar\n2311102210\nS1 IF-11-XX')
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Lecturer
    dos_heading = doc.add_heading('Dosen Pengampu :', 2)
    dos_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dos = doc.add_paragraph('Cahyo Prihantoro, S.Kom., M.Eng.')
    p_dos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # Institution
    inst_heading = doc.add_heading('LABORATORIUM HIGH PERFORMANCE\nFAKULTAS INFORMATIKA\nUNIVERSITAS TELKOM PURWOKERTO\n2026', 3)
    inst_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Content - Modul 13
    doc.add_heading('Dasar Praktikum', 1)
    doc.add_paragraph('Pada praktikum modul 13 ini, fokus pengembangan bergeser menuju eskalasi keamanan akses dan perancangan arsitektur database yang lebih kompleks pada proyek web Pixora...')
    
    doc.add_heading('Dasar Teori', 1)
    doc.add_heading('1.1 Manajemen Session', 2)
    doc.add_paragraph('Session adalah mekanisme penyimpanan data sementara di sisi server yang terikat pada interaksi pengguna tertentu.')
    
    doc.add_heading('1.2 Keamanan Berlapis via Middleware & Auth', 2)
    doc.add_paragraph('Middleware berfungsi sebagai pos pemeriksaan yang menyaring setiap HTTP Request yang masuk.')
    
    doc.add_heading('1.3 Model Relasi (Eloquent Relationships)', 2)
    doc.add_paragraph('Konsep One-to-Many diterapkan di Pixora; di mana satu Package dapat dipesan dan terkait dengan banyak data pemesanan (Booking).')
    
    doc.add_heading('PENGERJAAN & IMPLEMENTASI SISTEM', 1)
    doc.add_heading('2.1 Skema Autentikasi', 2)
    doc.add_paragraph('Akses ke menu pengelolaan paket admin kini dikunci sepenuhnya.')
    
    doc.add_heading('2.2 Relasi Entitas Database (One-to-Many)', 2)
    doc.add_paragraph('Tabel pendukung bookings dibuat dengan menjaga integritas data menggunakan foreignId.')
    
    doc.add_heading('3. Source Code Praktikum', 2)
    doc.add_paragraph('Source code lengkap dapat dilihat pada folder "Source Code/Modul 13".')
    
    doc.add_heading('HASIL TAMPILAN WEB (OUTPUT)', 2)
    try:
        doc.add_paragraph('1. Tampilan Halaman Login')
        doc.add_picture('/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/SS/login_modul_13.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[Gambar Tampilan Login]')
        
    try:
        doc.add_paragraph('2. Tampilan Header Auth di Layout Global')
        doc.add_picture('/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/SS/header_auth_modul_13.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[Gambar Tampilan Header Auth]')
        
    try:
        doc.add_paragraph('3. Tampilan Halaman Daftar Paket & Booking Pixora')
        doc.add_picture('/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/SS/product_list_modul_13.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[Gambar Tampilan Daftar Package dan Booking]')
        
    doc.add_page_break()
    
    # Tugas 8
    doc.add_heading('TUGAS PERTEMUAN 8', 1)
    doc.add_heading('1. Penjelasan tentang git branch', 2)
    doc.add_paragraph('Git branch adalah fitur dalam Git yang berfungsi menciptakan ruang kerja terpisah (cabang) dari repositori utama (main/master).')
    
    try:
        doc.add_paragraph('Membuat Git Branch dengan 2 akun berbeda (Untuk project Pixora):')
        doc.add_picture('/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/SS/git_branch.png', width=Inches(6.0))
    except Exception:
        pass
        
    doc.add_heading('2. Sistem Manajemen Studio Pixora', 2)
    doc.add_paragraph('Proyek ini dibangun menggunakan Laravel 11 dan ditujukan untuk mengimplementasikan manajemen basis data relasional pada layanan pemotretan studio fotography Pixora.')
    doc.add_paragraph('Source code lengkap dapat dilihat pada folder "Source Code/Tugas 8".')
    
    doc.add_heading('OUTPUT WEBSITE PIXORA (SS)', 2)
    images_tugas_8 = [
        ('1. Landing Page Pixora', 'landing_page_tugas_8.png'),
        ('2. Register Admin', 'register_tugas_8.png'),
        ('3. Login Admin', 'login_tugas_8.png'),
        ('4. Dashboard Admin Pixora', 'dashboard_tugas_8.png'),
        ('5. Tambah Paket Fotografi', 'tambah_data_tugas_8.png'),
        ('6. Edit Paket Fotografi', 'edit_data_tugas_8.png'),
    ]
    
    for title_img, img_file in images_tugas_8:
        doc.add_paragraph(title_img)
        try:
            doc.add_picture(f'/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/SS/{img_file}', width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f'[Gambar {title_img} missing or failed to load: {e}]')
    
    out_path = '/Users/haposansiregar/Downloads/abp/Haposan - 2311102210/Laporan_Praktikum_Modul_12_13.docx'
    doc.save(out_path)
    print(f"DOCX created at: {out_path}")

if __name__ == '__main__':
    main()
